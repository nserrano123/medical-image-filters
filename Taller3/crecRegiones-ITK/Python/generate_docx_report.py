"""
Generates a .docx report with embedded images for the segmentation comparison.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

REPORT_DIR = ".."
IMG_DIR = os.path.join(REPORT_DIR, "resultados_png")
OUTPUT = os.path.join(REPORT_DIR, "INFORME_Segmentacion.docx")

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)

def add_img(filename, width=6.0):
    path = os.path.join(IMG_DIR, filename)
    doc.add_picture(path, width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_obs(params, obs):
    p = doc.add_paragraph()
    r = p.add_run("Parámetros: ")
    r.bold = True
    p.add_run(params)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Observación: ")
    r2.bold = True
    p2.add_run(obs)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val

# ============================================================
# CONTENT
# ============================================================

add_heading("Informe: Segmentación por Crecimiento de Regiones con ITK", 0)

add_heading("1. Objetivo", 1)
doc.add_paragraph(
    "Segmentar los ventrículos cerebrales en un volumen de resonancia magnética (RM) "
    "ponderada en T1 utilizando dos técnicas de crecimiento de regiones implementadas "
    "con la librería ITK (Insight Toolkit):"
)
doc.add_paragraph("ConnectedThreshold — umbrales manuales definidos por el usuario.", style='List Bullet')
doc.add_paragraph("ConfidenceConnected — umbrales automáticos calculados estadísticamente.", style='List Bullet')
doc.add_paragraph(
    "Se busca comparar ambos métodos variando sus parámetros para evaluar la calidad "
    "de la segmentación y su sensibilidad a la configuración."
)

# ---
add_heading("2. Imagen de Entrada", 1)
doc.add_paragraph("Archivo: A1_grayT1.nii / A1_grayT1.nii.gz")
doc.add_paragraph("Tipo: Volumen 3D de resonancia magnética cerebral ponderada en T1, escala de grises.")
doc.add_paragraph("Dimensiones: 288 × 320 × 208 vóxeles.")
doc.add_paragraph(
    "Características T1: La sustancia blanca aparece clara (intensidad alta), la sustancia gris "
    "en tonos intermedios, y el líquido cefalorraquídeo (LCR) — incluyendo los ventrículos — "
    "aparece oscuro (intensidad baja). Esto hace que los ventrículos sean un buen candidato "
    "para segmentación por crecimiento de regiones, ya que presentan una intensidad relativamente "
    "homogénea y diferenciada del tejido circundante."
)
p = doc.add_paragraph()
r = p.add_run("Semilla utilizada en todos los experimentos: ")
r.bold = True
p.add_run("(132, 142, 96) — punto ubicado dentro del ventrículo lateral.")

# ---
add_heading("3. Descripción de los Métodos", 1)

add_heading("3.1 ConnectedThreshold (Umbrales Manuales)", 2)
doc.add_paragraph(
    "Este método parte de un punto semilla y examina los vóxeles vecinos. Si la intensidad "
    "de un vecino está dentro del rango [Lower, Upper] definido por el usuario, se agrega a "
    "la región segmentada. El proceso se repite hasta que no quedan más vecinos que cumplan el criterio."
)
add_table(
    ["Parámetro", "Descripción"],
    [
        ["LowerThreshold", "Umbral inferior de intensidad"],
        ["UpperThreshold", "Umbral superior de intensidad"],
        ["XSeed, YSeed, ZSeed", "Coordenadas de la semilla"],
    ]
)
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Ejemplo de ejecución: ")
r.bold = True
doc.add_paragraph("python3 ConnectedThreshold.py A1_grayT1.nii.gz output.nii.gz 100 170 132 142 96", style='No Spacing')
doc.add_paragraph()
doc.add_paragraph("Ventaja: Control total sobre el rango de intensidad.")
doc.add_paragraph("Desventaja: Requiere conocimiento previo de las intensidades de la estructura objetivo.")

add_heading("3.2 ConfidenceConnected (Umbrales Automáticos)", 2)
doc.add_paragraph(
    "También es crecimiento de regiones, pero los umbrales se calculan automáticamente. "
    "El algoritmo toma el vecindario alrededor de la semilla, calcula la media (μ) y "
    "desviación estándar (σ), y define el rango como [μ − f·σ, μ + f·σ], donde f es el "
    "multiplicador. Opcionalmente puede iterar: después de segmentar, recalcula μ y σ con "
    "los vóxeles ya incluidos y vuelve a crecer."
)
add_table(
    ["Parámetro", "Descripción"],
    [
        ["NumberOfIterations", "Veces que recalcula estadísticas y re-crece (0 = una pasada)"],
        ["Multiplier", "Factor f que multiplica σ. Mayor = más permisivo"],
        ["InitialNeighborhoodRadius", "Radio del vecindario para estadísticas iniciales (ej: 1 = cubo 3×3×3)"],
        ["XSeed, YSeed, ZSeed", "Coordenadas de la semilla"],
    ]
)
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Ejemplo de ejecución: ")
r.bold = True
doc.add_paragraph("python3 ConfidenceConnected.py A1_grayT1.nii.gz output.nii.gz 0 2 1 132 142 96", style='No Spacing')
doc.add_paragraph()
doc.add_paragraph("Ventaja: No requiere adivinar los umbrales; se adapta a la intensidad local.")
doc.add_paragraph("Desventaja: Depende de que la semilla esté en una zona representativa.")

# ---
add_heading("4. Experimentos y Resultados", 1)

add_heading("4.1 ConnectedThreshold — Variación de Umbrales", 2)

add_heading("Comparación general (corte axial)", 3)
add_img("comparacion_ConnectedThreshold.png")

add_heading("Prueba 1: Rango restrictivo — Lower=80, Upper=140", 3)
add_img("CT_narrow_80_140.png")
add_obs(
    "Lower=80, Upper=140",
    "La segmentación es muy conservadora. Solo captura una porción pequeña del ventrículo. "
    "El rango es demasiado estrecho y excluye vóxeles que sí pertenecen a la estructura. "
    "Archivo de salida: 37 KB."
)

add_heading("Prueba 2: Rango por defecto — Lower=100, Upper=170", 3)
add_img("CT_default_100_170.png")
add_obs(
    "Lower=100, Upper=170",
    "Buena segmentación del ventrículo. El rango captura la mayor parte de la estructura "
    "ventricular sin desbordarse significativamente. Este es el balance más adecuado de los cuatro. "
    "Archivo: 68 KB."
)

add_heading("Prueba 3: Rango amplio — Lower=60, Upper=200", 3)
add_img("CT_wide_60_200.png")
add_obs(
    "Lower=60, Upper=200",
    "La segmentación se desborda considerablemente, incluyendo tejido que no pertenece al "
    "ventrículo (sustancia gris, bordes). Archivo: 807 KB."
)

add_heading("Prueba 4: Rango muy amplio — Lower=50, Upper=250", 3)
add_img("CT_verywide_50_250.png")
add_obs(
    "Lower=50, Upper=250",
    "Desbordamiento severo. La segmentación invade gran parte del cerebro. Con un rango tan "
    "permisivo, prácticamente cualquier tejido es incluido. Archivo: 1.0 MB. Demuestra la "
    "importancia de elegir umbrales adecuados."
)

# ---
add_heading("4.2 ConfidenceConnected — Variación de Parámetros Estadísticos", 2)

add_heading("Comparación general (corte axial)", 3)
add_img("comparacion_ConfidenceConnected.png")

add_heading("Prueba 1: Conservador — Mult=1, Iter=0, Rad=1", 3)
add_img("CC_m1_i0_r1.png")
add_obs(
    "Multiplier=1, Iterations=0, Radius=1",
    "Segmentación conservadora. Con multiplicador 1 (rango μ ± 1σ), solo se incluyen vóxeles "
    "muy similares a la semilla. Captura el ventrículo de forma contenida. Archivo: 41 KB."
)

add_heading("Prueba 2: Default — Mult=2, Iter=0, Rad=1", 3)
add_img("CC_m2_i0_r1.png")
add_obs(
    "Multiplier=2, Iterations=0, Radius=1",
    "Con multiplicador 2 (μ ± 2σ) el rango es más amplio pero la estructura del ventrículo "
    "se mantiene bien delimitada. Buen punto de partida. Archivo: 41 KB."
)

add_heading("Prueba 3: Con iteraciones — Mult=2, Iter=3, Rad=1", 3)
add_img("CC_m2_i3_r1.png")
add_obs(
    "Multiplier=2, Iterations=3, Radius=1",
    "Al agregar 3 iteraciones, el algoritmo recalcula estadísticas y vuelve a crecer. El resultado "
    "es similar al anterior, sugiriendo que la región converge rápidamente. Archivo: 41 KB."
)

add_heading("Prueba 4: Agresivo — Mult=3, Iter=2, Rad=2", 3)
add_img("CC_m3_i2_r2.png")
add_obs(
    "Multiplier=3, Iterations=2, Radius=2",
    "Con multiplicador 3 y radio 2, la segmentación se desborda significativamente. Incluso el "
    "método automático puede fallar con parámetros agresivos. Archivo: 1.0 MB."
)

# ---
add_heading("5. Comparación de Métodos", 1)

add_table(
    ["Aspecto", "ConnectedThreshold", "ConfidenceConnected"],
    [
        ["Tipo de umbral", "Manual (usuario define Lower/Upper)", "Automático (calculado con μ ± f·σ)"],
        ["Facilidad de uso", "Requiere conocer las intensidades", "Solo requiere semilla y multiplicador"],
        ["Sensibilidad", "Muy sensible al rango elegido", "Más robusto con parámetros moderados"],
        ["Mejor resultado", "Lower=100, Upper=170", "Mult=2, Iter=0-3, Rad=1"],
        ["Riesgo de desbordamiento", "Alto con rangos amplios", "Menor con multiplicadores bajos"],
        ["Adaptabilidad", "No se adapta a variaciones locales", "Se adapta a la estadística local"],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Conclusión: ")
r.bold = True
p.add_run(
    "ConfidenceConnected es más robusto y fácil de usar para segmentar ventrículos, ya que "
    "calcula los umbrales automáticamente y es menos sensible a la elección de parámetros "
    "(siempre que el multiplicador sea razonable, entre 1 y 2.5). ConnectedThreshold ofrece "
    "más control pero requiere experimentación manual para encontrar el rango correcto, y es "
    "más propenso a desbordamiento con rangos amplios."
)

# ---
add_heading("6. Visor Interactivo", 1)
doc.add_paragraph(
    "Además de los scripts originales, se desarrolló un visor interactivo "
    "(interactive_segmentation.py) que permite:"
)
doc.add_paragraph("Visualizar el volumen original en cortes axial, coronal y sagital.", style='List Bullet')
doc.add_paragraph("Cambiar entre ambos métodos con radio buttons.", style='List Bullet')
doc.add_paragraph("Ajustar todos los parámetros en tiempo real con sliders.", style='List Bullet')
doc.add_paragraph("Ejecutar la segmentación y ver el resultado superpuesto sobre la imagen original.", style='List Bullet')
doc.add_paragraph("Visualizar la posición de la semilla con un marcador verde.", style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Ejecución: ")
r.bold = True
doc.add_paragraph("python3 interactive_segmentation.py A1_grayT1.nii.gz", style='No Spacing')

# ---
add_heading("7. Archivos del Proyecto", 1)
add_table(
    ["Archivo", "Descripción"],
    [
        ["A1_grayT1.nii.gz", "Volumen RM T1 de cerebro (entrada)"],
        ["ConnectedThreshold.py", "Segmentación con umbrales manuales"],
        ["ConfidenceConnected.py", "Segmentación con umbrales automáticos"],
        ["interactive_segmentation.py", "Visor interactivo con ambos métodos"],
        ["generate_comparison_images.py", "Generador de imágenes PNG comparativas"],
        ["resultados_png/", "Carpeta con todas las imágenes de resultados"],
    ]
)
doc.add_paragraph()
doc.add_paragraph("Requisitos: Python 3, ITK (pip install itk), matplotlib, numpy.")

# Save
doc.save(OUTPUT)
print(f"Report saved to: {OUTPUT}")
